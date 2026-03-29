"""Generate the CompressorIQ User Guide as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

doc.styles["Heading 1"].font.size = Pt(22)
doc.styles["Heading 2"].font.size = Pt(16)
doc.styles["Heading 3"].font.size = Pt(13)

# ── Cover ────────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("CompressorIQ")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("User Guide & Application Walkthrough")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

doc.add_paragraph()
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run("Compressor Service Intelligence Platform")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("Version 0.2.0  |  March 2026")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.add_page_break()

# ── Table of Contents placeholder ────────────────────────────────────────

doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Overview",
    "2. Application Structure",
    "3. Step 1 — Upload Maintenance Data",
    "4. Step 2 — Review the Fleet Dashboard",
    "5. Step 3 — Investigate a Specific Compressor",
    "6. Step 4 — Search and Explore Service Records",
    "7. Step 5 — Generate an AI Recommendation",
    "8. Step 6 — Follow the Prescribed Workflow",
    "9. Step 7 — Review All Workflows",
    "10. The Feedback Loop",
    "11. Typical Daily Scenarios",
    "12. Intelligence Engine — How It Works",
    "13. Data Pipeline — How Import Works",
    "14. Quick Reference",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ── 1. Overview ──────────────────────────────────────────────────────────

doc.add_heading("1. Overview", level=1)

doc.add_paragraph(
    "CompressorIQ is a compressor service intelligence platform that transforms "
    "unstructured maintenance records into actionable, evidence-based recommendations "
    "for field technicians and operations leaders."
)
doc.add_paragraph(
    "The platform ingests historical service data from spreadsheets, normalizes it "
    "into a structured database, and applies a 6-layer intelligence engine to generate "
    "prescriptive maintenance workflows — complete with confidence scores, similar "
    "historical cases, and step-by-step procedures."
)

doc.add_heading("What CompressorIQ Does", level=2)
bullets = [
    "Ingests and normalizes maintenance spreadsheets (SAP-style exports)",
    "Provides fleet-wide dashboards with KPIs, cost tracking, and issue trends",
    "Enables deep-dive into individual compressor history, issue frequency, and timelines",
    "Searches and filters across all service records with full-text search",
    "Generates AI-powered maintenance recommendations with explainable confidence",
    "Provides step-by-step prescribed workflows for field technicians",
    "Captures technician feedback to continuously improve future recommendations",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Who Uses It", level=2)

roles_data = [
    ("Field Technician", "Follows recommended workflows, captures feedback"),
    ("Maintenance Planner", "Reviews schedules, dispatches technicians, tracks outcomes"),
    ("Reliability Engineer", "Analyzes failure patterns, reviews issue frequency trends"),
    ("Operations Manager", "Monitors fleet health dashboards, reviews KPIs"),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Role"
hdr[1].text = "Primary Use"
for role, use in roles_data:
    row = table.add_row().cells
    row[0].text = role
    row[1].text = use

doc.add_paragraph()
doc.add_page_break()

# ── 2. Application Structure ─────────────────────────────────────────────

doc.add_heading("2. Application Structure", level=1)

doc.add_paragraph(
    "CompressorIQ has five main pages, accessible from the sidebar navigation on the left:"
)

pages_data = [
    ("Dashboard", "/", "Fleet-wide overview with KPIs, recent events, top issues, and machines needing attention"),
    ("Service Records", "/service-records", "Searchable, filterable list of all maintenance events with expandable detail"),
    ("Compressors", "/machines", "Asset list with per-compressor detail, issue frequency, and service timeline"),
    ("Workflows", "/workflow", "All AI-generated recommendations and prescribed maintenance workflows"),
    ("Upload Data", "/upload", "File upload interface for importing maintenance spreadsheets"),
]
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Page"
hdr[1].text = "URL Path"
hdr[2].text = "Purpose"
for page, path, purpose in pages_data:
    row = table.add_row().cells
    row[0].text = page
    row[1].text = path
    row[2].text = purpose

doc.add_paragraph()

doc.add_heading("Technology Stack", level=2)
stack = [
    ("Backend", "Python 3.12, FastAPI, SQLAlchemy 2.0, PostgreSQL / SQLite"),
    ("Frontend", "Next.js 16, React 19, TypeScript, Tailwind CSS 4"),
    ("Intelligence", "Rule-based inference, weighted similarity scoring, prescriptive workflows"),
    ("Data Pipeline", "pandas, openpyxl — 10-stage ingestion with full audit trail"),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Layer"
hdr[1].text = "Technologies"
for layer, tech in stack:
    row = table.add_row().cells
    row[0].text = layer
    row[1].text = tech

doc.add_paragraph()
doc.add_page_break()

# ── Step 1 ───────────────────────────────────────────────────────────────

doc.add_heading("3. Step 1 — Upload Maintenance Data", level=1)

p = doc.add_paragraph()
run = p.add_run("Who: ")
run.bold = True
p.add_run("Maintenance planner, data admin, or reliability engineer")
p = doc.add_paragraph()
run = p.add_run("Where: ")
run.bold = True
p.add_run("Sidebar > Upload Data")

doc.add_paragraph(
    "This is where everything starts. Before CompressorIQ can provide any intelligence, "
    "it needs historical maintenance records."
)

doc.add_heading("Procedure", level=2)
steps = [
    "Navigate to the Upload Data page from the sidebar.",
    "Drag and drop (or browse to select) a maintenance spreadsheet — .xlsx, .xls, or .csv.",
    'The file should follow the SAP-style export format with columns like "Plant", '
    '"Order & Description", "Customer Name", "Equipment", "Maintenance Activity Type", '
    '"Order Cost", "Run Hours", etc.',
    "Click Upload — the system runs a 10-stage ingestion pipeline behind the scenes.",
    "The Upload History table at the bottom shows the result — how many records were imported, "
    "whether it succeeded or failed, and any error messages.",
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_heading("What Happens Behind the Scenes", level=2)
pipeline_steps = [
    "Discovers and reads the workbook/CSV file",
    "Maps source columns to internal field names (case-insensitive matching)",
    "Normalizes values — activity types, event categories, action keywords",
    "Validates data quality — flags missing dates, negative costs, invalid codes",
    "Deduplicates against existing records via SHA-256 fingerprinting",
    "Persists normalized events, actions, notes, and measurements to the database",
    "Logs all data quality issues for audit review",
    "Generates a summary report of the import",
]
for step in pipeline_steps:
    doc.add_paragraph(step, style="List Bullet")

doc.add_heading("Current Dataset", level=2)
doc.add_paragraph(
    "The seed dataset is from Unit MC6068 — a single compressor with 303 service events "
    "spanning approximately 6 years of maintenance history (2020–2026). It includes 458 "
    "maintenance actions, 1,627 technician notes, and 339 measurements."
)

doc.add_heading("File Format Requirements", level=2)
doc.add_paragraph(
    "Files from any location on your computer can be uploaded. The system validates column "
    "compatibility before processing. If the file's columns don't match the expected SAP-style "
    "format, a clear error message will explain which columns were found versus which are expected."
)

doc.add_page_break()

# ── Step 2 ───────────────────────────────────────────────────────────────

doc.add_heading("4. Step 2 — Review the Fleet Dashboard", level=1)

p = doc.add_paragraph()
run = p.add_run("Who: ")
run.bold = True
p.add_run("Operations manager, reliability engineer, maintenance planner")
p = doc.add_paragraph()
run = p.add_run("Where: ")
run.bold = True
p.add_run("Sidebar > Dashboard")

doc.add_paragraph(
    "The dashboard is the operational home screen. After data is ingested, it immediately "
    "provides a fleet-wide overview."
)

doc.add_heading("KPI Cards (Top Row)", level=2)
kpis = [
    ("Total Service Events", "The complete count of maintenance records in the system."),
    ("Corrective Events", "Unplanned repairs — the ones you want to reduce."),
    ("Preventive Events", "Scheduled maintenance — the ones you want to optimize."),
    ("Average Cost", "Per-event cost across the fleet."),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "KPI"
hdr[1].text = "Description"
for kpi, desc in kpis:
    row = table.add_row().cells
    row[0].text = kpi
    row[1].text = desc

doc.add_paragraph()

doc.add_heading("Dashboard Sections", level=2)
sections = [
    ("Recent Service Events", "A table showing the latest maintenance activity. Each row is clickable "
     "and links to the Service Records page for deeper investigation."),
    ("Top Issue Categories", "A ranked breakdown showing what types of work dominate — corrective, "
     "oil sampling, emissions inspection, preventive maintenance, etc. Each category shows its count "
     "and percentage with a visual progress bar."),
    ("Machines Needing Attention", "Compressors with elevated recent corrective activity in the last "
     "30 days. This is the early warning system — units appearing here may need proactive intervention."),
]
for title, desc in sections:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ── Step 3 ───────────────────────────────────────────────────────────────

doc.add_heading("5. Step 3 — Investigate a Specific Compressor", level=1)

p = doc.add_paragraph()
run = p.add_run("Who: ")
run.bold = True
p.add_run("Reliability engineer, field service manager")
p = doc.add_paragraph()
run = p.add_run("Where: ")
run.bold = True
p.add_run("Sidebar > Compressors")

doc.add_heading("Layout", level=2)
doc.add_paragraph(
    "The page is split into two panels. The left panel shows the Asset List — all compressor "
    "units in the system. Clicking on a compressor loads its detail in the right panel."
)

doc.add_heading("Detail Panel Contents", level=2)
details = [
    ("Unit Header", "Unit ID (e.g., MC6068), equipment number, operational status, and current "
     "run hours."),
    ("Summary Cards", "Four cards showing total events, corrective count, preventive count, and "
     "the last service date for this unit."),
    ("Issue Frequency Table", "Every event category for this compressor — how many times it occurred, "
     "when it last happened, and the average run hours at occurrence. This tells you what keeps "
     "going wrong on this specific machine."),
    ("Service Timeline", "A chronological list of all service events, oldest to newest, with "
     "category badges. This tells you the full maintenance story of the machine."),
]
for title, desc in details:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph(
    'The compressor view answers: "Is this machine healthy? What are its recurring problems? '
    'When was it last serviced?"'
)

doc.add_page_break()

# ── Step 4 ───────────────────────────────────────────────────────────────

doc.add_heading("6. Step 4 — Search and Explore Service Records", level=1)

p = doc.add_paragraph()
run = p.add_run("Who: ")
run.bold = True
p.add_run("Field technician (pre-dispatch), maintenance planner, reliability engineer")
p = doc.add_paragraph()
run = p.add_run("Where: ")
run.bold = True
p.add_run("Sidebar > Service Records")

doc.add_paragraph(
    "This is the detailed event-level view with full search and filtering capabilities."
)

doc.add_heading("Search and Filter", level=2)
filters = [
    "Keyword search — searches across technician notes, order descriptions, and order numbers",
    "Category filter — dropdown of all event categories (corrective, preventive, oil sampling, etc.)",
    "Date range — From and To date pickers to narrow the time window",
]
for f in filters:
    doc.add_paragraph(f, style="List Bullet")

doc.add_heading("Event Table", level=2)
doc.add_paragraph(
    "The paginated table shows: Date, Order Number, Description, Category, Activity Type, "
    "Run Hours, and Cost. Results are paginated at 15 events per page with full navigation."
)

doc.add_heading("Expanded Detail View", level=2)
doc.add_paragraph("Click View on any row to expand it inline. The expanded view shows:")
expanded = [
    ("Technician Notes", "The full free-text notes from the field, cleaned and normalized. "
     "These often contain detailed descriptions of what was found, what was done, and any "
     "measurements taken."),
    ("Maintenance Actions", "Each discrete action extracted from the record — including the "
     "action type, component affected, technician name, date, and run hours at time of action."),
]
for title, desc in expanded:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph(
    "This is where a technician preparing for a dispatch would check: "
    '"What happened last time at this unit? What did the previous tech find?"'
)

doc.add_page_break()

# ── Step 5 ───────────────────────────────────────────────────────────────

doc.add_heading("7. Step 5 — Generate an AI Recommendation", level=1)

p = doc.add_paragraph()
run = p.add_run("Who: ")
run.bold = True
p.add_run("Field technician, maintenance planner")
p = doc.add_paragraph()
run = p.add_run("Where: ")
run.bold = True
p.add_run("Service Records page > Get Recommendation button")

doc.add_paragraph(
    "This is the core intelligence feature. For any service event, you can generate an "
    "AI-powered maintenance recommendation."
)

doc.add_heading("How to Generate", level=2)
gen_steps = [
    "On the Service Records page, find the event you're interested in.",
    'Click the "Get Recommendation" button on the row (or the larger button in the expanded detail view).',
    "CompressorIQ's 6-layer intelligence engine processes the event (typically 1–3 seconds).",
    "You are automatically redirected to the Workflow page for the generated recommendation.",
]
for i, step in enumerate(gen_steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_heading("What the Intelligence Engine Does", level=2)
layers = [
    ("Layer 1 — Analytics", "Analyzes action/issue frequencies, recurrence patterns, and time "
     "intervals between events for this machine."),
    ("Layer 2 — Rules Engine", "Infers the likely issue type from keywords, notes, event "
     "categories, and pattern matching against known failure modes."),
    ("Layer 3 — Similarity", "Finds similar historical cases using weighted scoring across: "
     "same machine, same category, keyword overlap (Jaccard similarity), and temporal recency."),
    ("Layer 4 — Workflow", "Generates a prescriptive step-by-step maintenance procedure "
     "tailored to the identified issue."),
    ("Layer 5 — Confidence", "Calculates a multi-factor confidence score from 6 auditable "
     "components. The score determines the confidence label: Low, Medium, or High."),
    ("Layer 5b — Explanation", "Writes a plain-language explanation where every sentence "
     "references actual data points from the machine's history."),
]
for title, desc in layers:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ── Step 6 ───────────────────────────────────────────────────────────────

doc.add_heading("8. Step 6 — Follow the Prescribed Workflow", level=1)

p = doc.add_paragraph()
run = p.add_run("Who: ")
run.bold = True
p.add_run("Field technician")
p = doc.add_paragraph()
run = p.add_run("Where: ")
run.bold = True
p.add_run("Sidebar > Workflows > click Open Workflow (or redirected from Step 5)")

doc.add_paragraph(
    "The workflow page is the technician's primary working screen during a maintenance event. "
    "It contains five sections:"
)

doc.add_heading("A. Recommendation Card", level=2)
doc.add_paragraph(
    "Displays the likely issue category, recommended action, confidence score with label "
    "(e.g., 77% — High), and a plain-language explanation. The explanation is evidence-based — "
    "every claim references actual data points from the machine's history."
)

doc.add_heading("B. Similar Historical Cases", level=2)
doc.add_paragraph(
    "A table of the most similar past events across the fleet. Each case shows its similarity "
    "score, what happened, and how it was resolved. This gives the technician context: "
    '"Other technicians faced this same pattern — here\'s what worked."'
)

doc.add_heading("C. Prescribed Workflow Steps", level=2)
doc.add_paragraph("A numbered checklist of specific actions to take. Each step includes:")
wf_items = [
    "The instruction itself — what exactly to do",
    "The rationale — why this step matters",
    "Required evidence — what to document or verify",
]
for item in wf_items:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph(
    "The technician checks off each step as they complete it and can add notes to any step "
    '(e.g., "Found 3/16 fitting cracked at stage 2 discharge").'
)

doc.add_heading("D. Technician Feedback Form", level=2)
doc.add_paragraph(
    "After completing the work, the technician records the outcome:"
)
feedback_items = [
    "What action was actually taken",
    "Whether the issue was resolved (yes/no)",
    "What parts were used",
    "Root cause found",
    "Any additional resolution notes",
]
for item in feedback_items:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph(
    "This feedback is critical — it feeds back into the intelligence engine so future "
    "recommendations for similar issues are weighted by what actually resolved them."
)

doc.add_heading("E. Status Controls", level=2)
statuses = [
    ("Accept", "Technician agrees with the recommendation and plans to follow it."),
    ("Mark Complete", "The work has been done and the workflow is finished."),
    ("Reject", "The recommendation was not applicable. This is a valuable learning signal "
     "for the system."),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Action"
hdr[1].text = "Meaning"
for action, meaning in statuses:
    row = table.add_row().cells
    row[0].text = action
    row[1].text = meaning

doc.add_paragraph()
doc.add_page_break()

# ── Step 7 ───────────────────────────────────────────────────────────────

doc.add_heading("9. Step 7 — Review All Workflows", level=1)

p = doc.add_paragraph()
run = p.add_run("Who: ")
run.bold = True
p.add_run("Maintenance planner, operations manager")
p = doc.add_paragraph()
run = p.add_run("Where: ")
run.bold = True
p.add_run("Sidebar > Workflows")

doc.add_paragraph(
    "The workflow index page shows all generated recommendations in one consolidated table:"
)

table_cols = [
    "Date generated",
    "Issue category identified",
    "Recommended action",
    "Confidence score and label",
    "Number of similar historical cases found",
    "Current status (Pending / Accepted / Completed / Rejected)",
]
for col in table_cols:
    doc.add_paragraph(col, style="List Bullet")

doc.add_paragraph(
    "This is the management view — track what's been recommended, what's been acted on, and "
    "what's still pending. It answers: Are technicians adopting the system's guidance? Are "
    "recommendations resolving issues?"
)

doc.add_page_break()

# ── Feedback Loop ────────────────────────────────────────────────────────

doc.add_heading("10. The Feedback Loop", level=1)

doc.add_paragraph("The real power of CompressorIQ is the closed loop:")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Upload Data  →  Analyze Patterns  →  Generate Recommendation  →  Technician Acts\n"
    "       ↑                                                                    |\n"
    "       └──────────────── Technician submits feedback ←──────────┘"
)
run.font.name = "Consolas"
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph(
    "Every time a technician submits feedback, the system gets smarter:"
)

feedback_benefits = [
    "Actions that resolved issues get boosted in future recommendations",
    "Rejected recommendations signal the rules engine to adjust its patterns",
    "Resolution rates per issue category improve over time",
    "Similar case matching becomes more accurate with more labeled outcomes",
    "Confidence scoring becomes better calibrated with real-world validation",
]
for b in feedback_benefits:
    doc.add_paragraph(b, style="List Bullet")

doc.add_page_break()

# ── Daily Scenarios ──────────────────────────────────────────────────────

doc.add_heading("11. Typical Daily Scenarios", level=1)

doc.add_heading("Morning Dispatch Planning", level=2)
p = doc.add_paragraph()
run = p.add_run("Role: ")
run.bold = True
p.add_run("Maintenance planner")

dispatch_steps = [
    "Open the Dashboard and scan for machines needing attention.",
    "Notice MC6068 is flagged with 3 corrective events in 30 days.",
    "Click through to the Compressors page and review its issue frequency — corrective events dominate.",
    "Open the latest service record to see what the previous technician noted.",
    'Click "Get Recommendation" to generate an AI-powered workflow.',
    "Share or print the workflow for the technician being dispatched.",
]
for i, step in enumerate(dispatch_steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_heading("Technician in the Field", level=2)
p = doc.add_paragraph()
run = p.add_run("Role: ")
run.bold = True
p.add_run("Field technician")

field_steps = [
    "Open Service Records and search for the unit you're dispatched to.",
    "Expand the most recent event to read the previous technician's notes.",
    'Click "Get Recommendation" for an AI-generated action plan.',
    "Follow the step-by-step workflow, checking off steps as you complete them.",
    "Add notes to relevant steps documenting what you find.",
    "Submit feedback when done — record what action you took, whether it resolved the issue, and what parts were used.",
]
for i, step in enumerate(field_steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_heading("Weekly Reliability Review", level=2)
p = doc.add_paragraph()
run = p.add_run("Role: ")
run.bold = True
p.add_run("Reliability engineer, operations manager")

review_steps = [
    "Start at the Dashboard — review top issue categories and corrective vs. preventive ratios.",
    "Open Compressors — compare machines with high corrective event counts.",
    "Drill into the worst performers to understand their issue frequency patterns.",
    "Open Workflows — review which recommendations were accepted vs. rejected.",
    "Identify systemic patterns: Are the same failure modes appearing across multiple machines?",
    "Use insights to adjust preventive maintenance schedules or plan component replacements.",
]
for i, step in enumerate(review_steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_page_break()

# ── Intelligence Engine ──────────────────────────────────────────────────

doc.add_heading("12. Intelligence Engine — How It Works", level=1)

doc.add_paragraph(
    "CompressorIQ's recommendation engine is a 6-layer stack. Each layer contributes "
    "a specific type of analysis, and the results are combined by the orchestrator into "
    "a single, explainable recommendation."
)

layers_detail = [
    ("Layer 1: Descriptive Analytics", [
        "Counts action and issue frequencies for the target machine",
        "Detects recurrence signals — repeated actions, escalation patterns, chronic issues",
        "Calculates time intervals between events",
        "Summarizes recent activity (last 30 and 90 days)",
    ]),
    ("Layer 2: Rules Engine", [
        "Infers the most likely issue category from keywords in notes, descriptions, and actions",
        "Uses configurable keyword/pattern matching rules",
        "Maps raw activity types to normalized categories",
    ]),
    ("Layer 3: Similarity Service", [
        "Finds the most similar historical events across the entire fleet",
        "Uses weighted multi-factor scoring: same machine (high weight), same category, "
        "keyword overlap via Jaccard similarity, and temporal recency decay",
        "Returns ranked similar cases with match explanations",
    ]),
    ("Layer 4: Workflow Service", [
        "Generates a step-by-step maintenance procedure based on the identified issue",
        "Each step includes instruction, rationale, and required evidence",
        "Adapts to the confidence level — low-confidence issues get triage-style workflows",
    ]),
    ("Layer 5: Confidence Service", [
        "Calculates a confidence score from 6 auditable factors",
        "Factors include: data volume, recurrence strength, similar case quality, rule match "
        "strength, resolution rate history, and temporal relevance",
        "Produces a label: Low (< 0.4), Medium (0.4–0.7), or High (> 0.7)",
    ]),
    ("Layer 5b: Explanation Service", [
        "Writes a plain-language explanation in 2–4 sentences",
        "Every sentence references actual data points — no vague generalities",
        "Designed to be readable by a field technician in under 30 seconds",
    ]),
]

for title, points in layers_detail:
    doc.add_heading(title, level=2)
    for point in points:
        doc.add_paragraph(point, style="List Bullet")

doc.add_page_break()

# ── Data Pipeline ────────────────────────────────────────────────────────

doc.add_heading("13. Data Pipeline — How Import Works", level=1)

doc.add_paragraph(
    "The ingestion pipeline processes uploaded files through 10 stages, ensuring data quality "
    "and full audit traceability."
)

pipeline_stages = [
    ("Stage 1: Discovery", "Identifies uploaded files and validates file type (.xlsx, .xls, .csv, .tsv)."),
    ("Stage 2: Workbook Read", "Opens the file, reads sheets, and extracts raw row data."),
    ("Stage 3: Column Mapping", "Maps source column names to internal field names. Supports "
     "case-insensitive and whitespace-normalized matching."),
    ("Stage 4: Raw Preservation", "Stores every source row exactly as read, as JSON, for full "
     "auditability."),
    ("Stage 5: Normalization", "Transforms raw values — normalizes activity types, categorizes "
     "events, extracts actions from free text, parses dates and numbers."),
    ("Stage 6: Validation", "Applies business rules — checks for missing dates, negative costs, "
     "invalid codes. Logs issues at info/warning/error severity levels."),
    ("Stage 7: Deduplication", "Computes SHA-256 fingerprints of file contents and row data. "
     "Skips rows that already exist in the database."),
    ("Stage 8: Persistence", "Writes normalized events, actions, notes, measurements, and master "
     "data (compressors, sites, technicians) to the database."),
    ("Stage 9: Issue Logging", "Persists all data quality issues found during import for later review."),
    ("Stage 10: Report", "Generates a summary: rows scanned, imported, skipped, errored, plus "
     "counts of created entities."),
]

for title, desc in pipeline_stages:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()

doc.add_heading("Data Model", level=2)
doc.add_paragraph(
    "The database contains 23 tables organized into 4 zones:"
)
zones = [
    ("Import Zone", "import_batches, import_files, import_sheets, raw_service_rows, import_issue_log"),
    ("Master Zone", "compressors, sites, technicians, issue_categories, action_types"),
    ("Event Zone", "service_events, service_event_actions, service_event_notes, service_event_measurements"),
    ("Analytics Zone", "recommendations, similar_cases, workflow_steps, feedback"),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Zone"
hdr[1].text = "Tables"
for zone, tables_list in zones:
    row = table.add_row().cells
    row[0].text = zone
    row[1].text = tables_list

doc.add_paragraph()
doc.add_page_break()

# ── Quick Reference ──────────────────────────────────────────────────────

doc.add_heading("14. Quick Reference", level=1)

doc.add_heading("URLs", level=2)
urls = [
    ("Frontend (Dashboard)", "http://localhost:3000"),
    ("Backend API", "http://127.0.0.1:8001"),
    ("API Documentation", "http://127.0.0.1:8001/docs"),
    ("Health Check", "http://127.0.0.1:8001/health"),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Service"
hdr[1].text = "URL"
for service, url in urls:
    row = table.add_row().cells
    row[0].text = service
    row[1].text = url

doc.add_paragraph()

doc.add_heading("Key API Endpoints", level=2)
endpoints = [
    ("GET", "/api/dashboard/summary", "Fleet overview KPIs"),
    ("GET", "/api/service-events/", "Paginated event list with filters"),
    ("GET", "/api/service-events/{id}", "Full event detail with actions/notes"),
    ("GET", "/api/compressors/", "List all compressor assets"),
    ("GET", "/api/compressors/{id}", "Compressor detail with stats"),
    ("GET", "/api/recommendations/", "List all recommendations"),
    ("POST", "/api/recommendations/generate/{event_id}", "Generate recommendation"),
    ("POST", "/api/ingestion/upload", "Upload and import a file"),
    ("GET", "/api/ingestion/uploads", "Upload history"),
    ("POST", "/api/feedback/", "Submit technician feedback"),
]
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Method"
hdr[1].text = "Endpoint"
hdr[2].text = "Description"
for method, endpoint, desc in endpoints:
    row = table.add_row().cells
    row[0].text = method
    row[1].text = endpoint
    row[2].text = desc

doc.add_paragraph()

doc.add_heading("Supported File Formats", level=2)
formats = [".xlsx (Excel 2007+)", ".xls (Legacy Excel)", ".csv (Comma-separated)", ".tsv (Tab-separated)"]
for f in formats:
    doc.add_paragraph(f, style="List Bullet")

doc.add_paragraph()
doc.add_paragraph()

# ── Footer ───────────────────────────────────────────────────────────────

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("— End of Document —")
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
run.font.size = Pt(10)
run.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CompressorIQ v0.2.0 — Compressor Service Intelligence Platform")
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
run.font.size = Pt(9)

# ── Save ─────────────────────────────────────────────────────────────────

output_path = os.path.join(os.path.dirname(__file__), "CompressorIQ_User_Guide.docx")
doc.save(output_path)
print(f"Saved to: {output_path}")
